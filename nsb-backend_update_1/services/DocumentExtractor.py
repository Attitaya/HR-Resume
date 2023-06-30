import re
import docx
from docx.enum.text import WD_BREAK
class DocumentExtractor():
    def extract(self,path):
        data_dict= {}
        text = docx.Document(path)
        for i, para in enumerate(text.paragraphs):
            # Check if the paragraph is in bold and has underline
            if para.runs and (para.runs[0].bold or para.runs[0].underline):
                # Get the heading text
                heading_text = para.text.strip()

                # If there are multiple lines in the heading, select the shortest one
                if '\n' in heading_text:
                    heading_text = min(heading_text.split('\n'), key=len)

                # reprocess to get the real topic
                pattern = r'\s{2,}|[“]| โดย| ประกอบด้วย|[.]{2,}|…|	|\t| นำโดย| ด้วย|ของ '
                real = re.split(pattern, heading_text)
                for n in range (len(real)):
                    if len(real[n]) >= 1:
                        header = real[n].strip()
                        lastvalue = int(n)
                        break

                # the non-topic words in heading_text is considered a content of the topics in the first line
                value_fix = ""
                for words in real[lastvalue+1:]:
                    value_fix += words+" "

                # Initialize the value as an empty string
                value_text = "" + value_fix

                # Loop through the subsequent paragraphs until the next heading is encountered
                j = i + 1
                while j < len(text.paragraphs):
                    next_para = text.paragraphs[j]
                    # Check if the next paragraph is in bold and has underline, or is in bold only
                    if (next_para.runs and next_para.runs[0].bold and next_para.runs[0].underline) or (next_para.runs and next_para.runs[0].bold):
                        break
                    else:
                        # Add the text to the value string
                        value_text += next_para.text.strip()
                        # Add a separator between paragraphs
                        if next_para.runs and (next_para.runs[0].text == "\n" or next_para.runs[0].text == WD_BREAK.PAGE):
                            value_text += "\n\n"
                        else:
                            value_text += " "
                    j += 1
                # Add the heading and value to the dictionary
                data_dict[header] = value_text.strip()
        return data_dict

################################# For batch test #################################